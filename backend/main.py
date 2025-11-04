from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime
from pymongo import MongoClient
from fastapi.responses import JSONResponse
from io import BytesIO
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor
import pytz
from dateutil import parser
from xml.sax.saxutils import escape
from urllib.parse import urlparse, urlunparse

# === MongoDB setup ===
# client = MongoClient("mongodb://10.226.49.29:27017/")
client = MongoClient("mongodb://cdac:user%40cdac!@13.201.47.59:27017/?authSource=csmart")
# db = client["link_monitor"]
db = client["csmart"]
links_col = db["ai_video_links"]
logs_col = db["download_logs"]

app = FastAPI(title="IT Rules Link Logger API")

# === CORS Middleware ===
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

IST = pytz.timezone("Asia/Kolkata")

# === Pydantic Models ===
class LinkModel(BaseModel):
    url: str
    comments: str | None = None

class LogModel(BaseModel):
    from_date: str
    to_date: str
    count: int
    user: str


# === Utility ===
def clean_instagram_url(url: str) -> str:
    parsed = urlparse(url)
    clean_url = urlunparse((parsed.scheme, parsed.netloc, parsed.path, "", "", ""))
    if not clean_url.endswith("/"):
        clean_url += "/"
    return clean_url


def detect_platform(url: str) -> str:
    url = url.lower()
    if "twitter.com" in url or "x.com" in url: return "Twitter"
    if "facebook.com" in url: return "Facebook"
    if "instagram.com" in url: return "Instagram"
    if "youtube.com" in url: return "YouTube"
    if "t.me" in url or "telegram.org" in url: return "Telegram"
    if "whatsapp.com" in url: return "WhatsApp"
    if "reddit.com" in url: return "Reddit"
    return "Other"


# === API Endpoints ===

@app.post("/add_link/")
async def add_link(link: LinkModel):
    try:
        now_ist = datetime.now(IST)
        now_utc = now_ist.astimezone(pytz.UTC)  # ✅ store in UTC
        platform = detect_platform(link.url.strip())

        url = link.url.strip()
        if platform == "Instagram":
            url = clean_instagram_url(url)

        existing = links_col.find_one({"url": url})
        if existing:
            return JSONResponse(
                content={"message": "⚠️ This link already exists in the database.", "platform": platform},
                status_code=409
            )

        record = {
            "url": url,
            "platform": platform,
            "comments": link.comments,
            "rule_violation": "3(1)(b) (ii, v)",
            "action_status": "Not Taken Down",
            "timestamp": now_utc,  # ✅ UTC datetime object
        }

        links_col.insert_one(record)
        return {"message": "✅ Link added successfully", "platform": platform}
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


@app.get("/get_links/")
async def get_links(from_date: str, to_date: str):
    try:
        # Convert frontend IST timestamps to UTC for DB filter
        from_dt = parser.isoparse(from_date).astimezone(pytz.UTC)
        to_dt = parser.isoparse(to_date).astimezone(pytz.UTC)

        data = list(
            links_col.find({"timestamp": {"$gte": from_dt, "$lte": to_dt}}, {"_id": 0})
        )

        # Convert for display
        for item in data:
            if "timestamp" in item and isinstance(item["timestamp"], datetime):
                item["timestamp"] = (
                    item["timestamp"]
                    .astimezone(IST)
                    .strftime("%d %b %Y, %I:%M %p")  # AM/PM format
                )

        return {"data": data}
    except Exception as e:
        return {"error": str(e)}


@app.post("/log_download/")
async def log_download(log: LogModel):
    try:
        from_dt_ist = parser.isoparse(log.from_date)
        to_dt_ist = parser.isoparse(log.to_date)
        # ✅ Convert to UTC before storing
        from_dt_utc = from_dt_ist.astimezone(pytz.UTC)
        to_dt_utc = to_dt_ist.astimezone(pytz.UTC)
        now_utc = datetime.now(pytz.UTC)

        logs_col.insert_one({
            "from_date": from_dt_utc,
            "to_date": to_dt_utc,
            "count": log.count,
            "user": log.user,
            "timestamp": now_utc
        })

        return {"message": "✅ Download log saved successfully."}
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


@app.get("/get_logs/")
async def get_logs():
    try:
        logs = list(logs_col.find({}, {"_id": 0}).sort("timestamp", -1))
        formatted_logs = []

        for log in logs:
            for key in ["from_date", "to_date", "timestamp"]:
                if key in log and isinstance(log[key], datetime):
                    log[key] = (
                        log[key]
                        .astimezone(IST)
                        .strftime("%d %b %Y, %I:%M %p")  # ✅ IST AM/PM
                    )
            formatted_logs.append(log)

        return {"logs": formatted_logs}
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


# === EXPORT SECTION ===
@app.get("/export/")
async def export(from_date: str, to_date: str, file_type: str = "pdf"):
    try:
        from_dt = parser.isoparse(from_date)
        to_dt = parser.isoparse(to_date)
        data = list(links_col.find(
            {"timestamp": {"$gte": from_dt, "$lte": to_dt}},
            {"_id": 0}
        ))
    except Exception as e:
        return JSONResponse(content={"error": f"Invalid date: {str(e)}"}, status_code=400)

    if not data:
        return JSONResponse(content={"error": "No records found."}, status_code=404)
    
    # ✅ Convert UTC → IST for all datetime fields before building report
    IST = pytz.timezone("Asia/Kolkata")
    for record in data:
        if "timestamp" in record and isinstance(record["timestamp"], datetime):
            record["timestamp"] = record["timestamp"].astimezone(IST).strftime("%d %b %Y, %I:%M %p")


    df = pd.DataFrame(data)
    grouped = df.groupby("platform")

    # === PDF Export ===
    if file_type == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            rightMargin=40, leftMargin=40,
            topMargin=60, bottomMargin=40
        )
        elements = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            name="title_style", alignment=TA_CENTER,
            fontSize=16, textColor=colors.HexColor("#002147"),
            leading=20
        )
        subtitle_style = ParagraphStyle(
            name="subtitle_style", alignment=TA_CENTER,
            fontSize=10, leading=14, textColor=colors.black
        )
        platform_style = ParagraphStyle(
            name="platform_style", fontSize=12,
            textColor=colors.HexColor("#002147"),
            alignment=TA_LEFT, leftIndent=8,
            spaceBefore=8, spaceAfter=8
        )
        url_style = ParagraphStyle(
            name="url_style", alignment=TA_LEFT,
            textColor=colors.HexColor("#0000EE"),
            underline=True, fontSize=9
        )

        # Header
        elements.append(Paragraph("<b>Deepfake or Manipulated Content Report</b>", title_style))
        elements.append(Paragraph("<b>Related to Hon’ble PM Modi – Actionable Report</b>", title_style))
        elements.append(Paragraph(f"({datetime.now(IST).strftime('%d %B, %Y')})", subtitle_style))
        elements.append(Spacer(1, 15))

        for platform, group in grouped:
            elements.append(Paragraph(f"<b>{platform}</b>", platform_style))
            elements.append(Spacer(1, 6))

            header_style = ParagraphStyle(
                name="header_style",
                parent=styles["Normal"],
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
                fontSize=9,
                leading=12,
            )

            header_text = "<b>Relevant Violation<br/>of IT Rules, 2021</b>"
            table_data = [["S.No", "URL", Paragraph(header_text, header_style), "Action Status", "Comments"]]

            for i, row in enumerate(group.itertuples(), 1):
                url_para = Paragraph(
                    f"<a href='{escape(row.url)}' color='blue'><u>{escape(row.url)}</u></a>", url_style
                )
                table_data.append([
                    str(i), url_para,
                    row.rule_violation,
                    row.action_status,
                    row.comments or ""
                ])

            # fixed, consistent widths
            t = Table(table_data, colWidths=[35, 180, 115, 100, 80])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#dce6f1")),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                # Padding for better row spacing
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fb")]),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 14))

        doc.build(elements)
        buffer.seek(0)
        return JSONResponse(
            content={"file": buffer.getvalue().decode("latin1")},
            headers={"Content-Disposition": f"attachment; filename=violations_{from_date}_{to_date}.pdf"},
        )

    # === DOCX Export ===
    elif file_type == "docx":
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE

        def set_cell_padding(cell, top=120, bottom=120, left=80, right=80):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side, val in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)

        def add_hyperlink(paragraph, text, url):
            part = paragraph.part
            r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0000FF")
            rPr.append(color)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
            new_run.append(rPr)
            text_elem = OxmlElement("w:t")
            text_elem.text = text
            new_run.append(text_elem)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        doc = Document()
        for section in doc.sections:
            section.top_margin = Pt(40)
            section.bottom_margin = Pt(40)
            section.left_margin = Pt(40)
            section.right_margin = Pt(40)

        # Header
        title = doc.add_paragraph("Deepfake or Manipulated Content Report")
        title.alignment = 1
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 33, 71)

        subtitle = doc.add_paragraph("Related to Hon’ble PM Modi – Actionable Report")
        subtitle.alignment = 1
        run = subtitle.runs[0]
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 33, 71)

        date_para = doc.add_paragraph(f"({datetime.now(IST).strftime('%d %B, %Y')})")
        date_para.alignment = 1
        doc.add_paragraph()

        for platform, group in grouped:
            p = doc.add_paragraph()
            run = p.add_run(platform)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 33, 71)
            p.paragraph_format.space_after = Pt(6)

            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.autofit = False

            # Match PDF widths
            widths = [Inches(0.6), Inches(2.5), Inches(1.6), Inches(1.8), Inches(1.3)]
            for i, width in enumerate(widths):
                table.columns[i].width = width

            hdr_cells = table.rows[0].cells
            headers = ["S.No", "URL", "Relevant Violation\nof IT Rules, 2021", "Action Status", "Comments"]
            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)
                hdr_cells[i].paragraphs[0].alignment = 1

            for i, row in enumerate(group.itertuples(), 1):
                cells = table.add_row().cells
                cells[0].text = str(i)
                add_hyperlink(cells[1].paragraphs[0], row.url, row.url)
                cells[2].text = row.rule_violation
                cells[3].text = row.action_status
                cells[4].text = row.comments or ""

                for cell in cells:
                    set_cell_padding(cell)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(3)
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph()

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return JSONResponse(
            content={"file": buf.getvalue().decode("latin1")},
            headers={"Content-Disposition": f"attachment; filename=violations_{from_date}_{to_date}.docx"},
        )




@app.get("/")
async def home():
    return {"message": "✅ IT Rules Logger API is running"}
